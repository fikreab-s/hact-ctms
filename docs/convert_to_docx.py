"""
Convert HACT E2E System Workflow Architecture markdown to DOCX.
Produces a professional document suitable for pod lead review.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

MD_FILE = Path(r"C:\Users\hello\Desktop\HACT project\docs\HACT_E2E_System_Workflow_Architecture.md")
OUT_FILE = MD_FILE.with_suffix(".docx")

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# ── Define styles ──
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for i in range(1, 4):
    hs = doc.styles[f"Heading {i}"]
    hs.font.name = "Calibri"
    hs.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    if i == 1:
        hs.font.size = Pt(20)
        hs.font.bold = True
    elif i == 2:
        hs.font.size = Pt(15)
        hs.font.bold = True
    else:
        hs.font.size = Pt(12)
        hs.font.bold = True


def add_code_block(doc, text):
    """Add a monospaced code block paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)


def add_table_from_rows(doc, headers, rows):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val.strip()
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()  # spacer


def parse_table(lines, start_idx):
    """Parse a markdown table starting at start_idx. Returns (headers, rows, end_idx)."""
    headers = [c.strip() for c in lines[start_idx].strip().strip("|").split("|")]
    # Skip separator line
    rows = []
    i = start_idx + 2
    while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
        row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        rows.append(row)
        i += 1
    return headers, rows, i


# ── Read and parse markdown ──
content = MD_FILE.read_text(encoding="utf-8")
lines = content.split("\n")

# ── Cover page ──
doc.add_paragraph()
doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("HACT CTMS")
title_run.font.size = Pt(36)
title_run.bold = True
title_run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subtitle_p.add_run("End-to-End System Workflow & Architecture")
sub_run.font.size = Pt(18)
sub_run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

doc.add_paragraph()

desc_p = doc.add_paragraph()
desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
desc_run = desc_p.add_run("Persona-Based Operational Guide\nIncluding Mobile EDC")
desc_run.font.size = Pt(14)
desc_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_paragraph()
doc.add_paragraph()

meta_items = [
    ("Version", "2.0 — June 2026"),
    ("Prepared for", "Pod Lead Review"),
    ("Project", "HACT Clinical Trial Management System"),
    ("Protocol", "PSBI Neonatal Sepsis Trial (PSBI-2026-001)"),
    ("Compliance", "ICH E6(R2/R3), 21 CFR Part 11, FDA, ALCOA+"),
]
meta_table = doc.add_table(rows=len(meta_items), cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (k, v) in enumerate(meta_items):
    meta_table.rows[i].cells[0].text = k
    meta_table.rows[i].cells[1].text = v
    for p in meta_table.rows[i].cells[0].paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)
    for p in meta_table.rows[i].cells[1].paragraphs:
        for run in p.runs:
            run.font.size = Pt(10)

doc.add_page_break()

# ── Parse body content ──
i = 0
in_code_block = False
code_buffer = []

while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    # Skip front matter lines (already on cover page)
    if stripped.startswith("> **") or stripped == ">" or stripped == "---":
        i += 1
        continue

    # Code blocks
    if stripped.startswith("```"):
        if in_code_block:
            # End code block
            add_code_block(doc, "\n".join(code_buffer))
            code_buffer = []
            in_code_block = False
        else:
            in_code_block = True
        i += 1
        continue

    if in_code_block:
        code_buffer.append(line.rstrip())
        i += 1
        continue

    # Headings
    if stripped.startswith("# ") and not stripped.startswith("##"):
        # H1 — main title (skip, already on cover)
        i += 1
        continue
    elif stripped.startswith("## "):
        heading_text = stripped.lstrip("#").strip()
        # Clean markdown formatting
        heading_text = re.sub(r'[&]amp;', '&', heading_text)
        doc.add_heading(heading_text, level=1)
        i += 1
        continue
    elif stripped.startswith("### "):
        heading_text = stripped.lstrip("#").strip()
        heading_text = re.sub(r'[&]amp;', '&', heading_text)
        doc.add_heading(heading_text, level=2)
        i += 1
        continue
    elif stripped.startswith("#### "):
        heading_text = stripped.lstrip("#").strip()
        heading_text = re.sub(r'[&]amp;', '&', heading_text)
        doc.add_heading(heading_text, level=3)
        i += 1
        continue

    # Tables
    if "|" in stripped and stripped.startswith("|") and i + 1 < len(lines):
        next_line = lines[i + 1].strip() if i + 1 < len(lines) else ""
        if "|" in next_line and re.match(r"^\|[\s\-:|]+\|", next_line):
            headers, rows, end_idx = parse_table(lines, i)
            if headers and rows:
                add_table_from_rows(doc, headers, rows)
            i = end_idx
            continue

    # Bullet points
    if stripped.startswith("- ") or stripped.startswith("* "):
        bullet_text = stripped[2:].strip()
        # Clean HTML entities
        bullet_text = re.sub(r'&amp;', '&', bullet_text)
        bullet_text = re.sub(r'&lt;', '<', bullet_text)
        bullet_text = re.sub(r'&gt;', '>', bullet_text)
        # Remove markdown bold/italic
        bullet_text = re.sub(r'\*\*(.+?)\*\*', r'\1', bullet_text)
        bullet_text = re.sub(r'\*(.+?)\*', r'\1', bullet_text)
        p = doc.add_paragraph(bullet_text, style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(9)
        i += 1
        continue

    # Regular text (skip empty lines)
    if stripped:
        # Clean markdown formatting
        text = stripped
        text = re.sub(r'&amp;', '&', text)
        text = re.sub(r'&lt;', '<', text)
        text = re.sub(r'&gt;', '>', text)

        # Handle bold text with runs
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.+?\*\*)', text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
                run.font.size = Pt(10)
            else:
                run = p.add_run(part)
                run.font.size = Pt(10)
        i += 1
        continue

    i += 1

# ── Footer ──
doc.add_paragraph()
doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_p.add_run("Document prepared for HACT CTMS Pod Lead review — June 2026")
footer_run.font.size = Pt(9)
footer_run.italic = True
footer_run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

# ── Save ──
doc.save(str(OUT_FILE))
print(f"✅ DOCX created: {OUT_FILE}")
print(f"   Size: {OUT_FILE.stat().st_size / 1024:.1f} KB")
print(f"   Pages: ~{len(doc.paragraphs) // 40 + 1} (estimated)")
print()
print("To convert to PDF:")
print("  1. Open the .docx file in Microsoft Word")
print("  2. File → Save As → Choose PDF format")
print("  Or: File → Export → Create PDF/XPS")
